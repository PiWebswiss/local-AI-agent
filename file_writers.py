# Enable postponed evaluation of type annotations.
from __future__ import annotations

# Regex utilities are used for default "should transform" filtering.
import re
# Path is used for safe filesystem path handling.
from pathlib import Path
# Generic type hints for callback signatures and docx objects.
from typing import Any, Callable


# Custom exception for all write/transform failures from this module.
class FileWriteError(RuntimeError):
    # No extra behavior; this class exists for clearer error handling upstream.
    pass


# Ensure the parent directory exists before writing output files.
def _ensure_parent_dir(path: Path) -> None:
    # Try to create parent directories (no error if they already exist).
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
    # Convert low-level OS error into module-specific error type.
    except OSError as e:
        raise FileWriteError(f"Failed to create output directory: {path.parent} ({e})") from e


# Write plain UTF-8 text to a file path.
def write_text_file(path: str | Path, text: str) -> None:
    # Normalize to Path object.
    p = Path(path)
    # Ensure output folder exists.
    _ensure_parent_dir(p)
    # Attempt the write operation.
    try:
        p.write_text(text, encoding="utf-8")
    # Map filesystem errors to FileWriteError.
    except OSError as e:
        raise FileWriteError(f"Failed to write file: {p} ({e})") from e


# Write text into a new DOCX file (one line -> one paragraph).
def write_docx_file(path: str | Path, text: str) -> None:
    try:
        from docx import Document  # type: ignore
    # Surface a clear dependency message if import fails.
    except Exception as e:  # pragma: no cover
        raise FileWriteError(f"DOCX output requires 'python-docx'. ({e})") from e

    # Normalize destination path.
    p = Path(path)
    # Ensure destination directory exists.
    _ensure_parent_dir(p)

    # Create a new empty DOCX document.
    doc = Document()
    # Preserve input line breaks by adding one paragraph per line.
    # Note: this does not preserve existing complex formatting (runs/styles/tables).
    for line in (text or "").splitlines():
        doc.add_paragraph(line)

    # Save document to disk.
    try:
        doc.save(str(p))
    # Map write failures to FileWriteError.
    except OSError as e:
        raise FileWriteError(f"Failed to write DOCX: {p} ({e})") from e


# Add a paragraph once by tracking its underlying XML node id.
def _add_paragraph_once(paragraph: Any, out: list[Any], seen: set[int]) -> None:
    # Build identity key from internal paragraph XML object.
    key = id(paragraph._p)
    # Skip duplicates.
    if key in seen:
        return
    # Mark paragraph as seen.
    seen.add(key)
    # Append paragraph to output list.
    out.append(paragraph)


# Recursively collect paragraphs contained in a table.
def _collect_table_paragraphs(table: Any, out: list[Any], seen: set[int]) -> None:
    # Iterate every row in table.
    for row in table.rows:
        # Iterate every cell in row.
        for cell in row.cells:
            # Recurse through each cell's paragraphs/tables.
            _collect_cell_paragraphs(cell, out, seen)


# Collect all paragraphs from a cell, including nested tables.
def _collect_cell_paragraphs(cell: Any, out: list[Any], seen: set[int]) -> None:
    # Add direct paragraphs in this cell.
    for paragraph in cell.paragraphs:
        _add_paragraph_once(paragraph, out, seen)
    # Recurse into nested tables.
    for table in cell.tables:
        _collect_table_paragraphs(table, out, seen)


# Collect all paragraphs from body, tables, headers, and footers.
def _collect_docx_paragraphs(doc: Any) -> list[Any]:
    # Output paragraph list.
    out: list[Any] = []
    # Set of seen paragraph identities to avoid duplicates.
    seen: set[int] = set()

    # Collect body paragraphs.
    for paragraph in doc.paragraphs:
        _add_paragraph_once(paragraph, out, seen)
    # Collect body table paragraphs.
    for table in doc.tables:
        _collect_table_paragraphs(table, out, seen)

    # Collect paragraphs from all section header/footer variants.
    for section in doc.sections:
        # Iterate each header/footer attribute name.
        for attr in (
            "header",
            "first_page_header",
            "even_page_header",
            "footer",
            "first_page_footer",
            "even_page_footer",
        ):
            # Resolve part object from section.
            part = getattr(section, attr, None)
            # Skip missing parts.
            if part is None:
                continue
            # Collect direct paragraphs from the part.
            for paragraph in part.paragraphs:
                _add_paragraph_once(paragraph, out, seen)
            # Collect paragraphs inside tables in the part.
            for table in part.tables:
                _collect_table_paragraphs(table, out, seen)

    # Return deduplicated paragraph list.
    return out


# Replace paragraph text while preserving run structure as much as possible.
def _set_paragraph_text_preserving_runs(paragraph: Any, new_text: str) -> None:
    # Snapshot paragraph runs.
    runs = list(paragraph.runs)
    # If paragraph has no runs, create one.
    if not runs:
        paragraph.add_run(new_text)
        return

    # Measure original run text lengths.
    old_lengths = [len(run.text or "") for run in runs]
    # Total original characters across all runs.
    old_total = sum(old_lengths)
    # Target character length.
    new_len = len(new_text)

    # If new text is empty, clear all runs.
    if new_len == 0:
        for run in runs:
            run.text = ""
        return

    # If original had no textual width, put all text in first run.
    if old_total <= 0:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
        return

    # Build proportional cut offsets into new_text.
    cuts: list[int] = [0]
    # Running total of original text consumed.
    cumulative = 0
    # Iterate all runs except the last to compute split points.
    for length in old_lengths[:-1]:
        cumulative += length
        cut = round((cumulative / old_total) * new_len)
        # Enforce non-decreasing offsets.
        if cut < cuts[-1]:
            cut = cuts[-1]
        # Clamp to target length.
        if cut > new_len:
            cut = new_len
        # Store cut offset.
        cuts.append(cut)
    # Add final endpoint.
    cuts.append(new_len)

    # Rewrite each run with its mapped slice.
    for idx, run in enumerate(runs):
        run.text = new_text[cuts[idx] : cuts[idx + 1]]


# Default rule: transform non-empty lines that contain alphabetic letters.
def _should_transform_default(text: str) -> bool:
    # Trim whitespace.
    s = (text or "").strip()
    # Skip blank values.
    if not s:
        return False
    # Return true if line contains letters (including accented ranges).
    return bool(re.search(r"[A-Za-zÀ-ÖØ-öø-ÿ]", s))


# Transform an existing DOCX and save output while preserving run formatting.
def transform_docx_preserving_format(
    input_path: str | Path,
    output_path: str | Path,
    *,
    transform: Callable[[str], str],
    should_transform: Callable[[str], bool] | None = None,
    on_progress: Callable[[int, int], None] | None = None,
) -> tuple[int, int]:
    try:
        from docx import Document  # type: ignore
    # Report missing dependency as FileWriteError.
    except Exception as e:  # pragma: no cover
        raise FileWriteError(f"DOCX support requires 'python-docx'. ({e})") from e

    # Normalize source path.
    src = Path(input_path)
    # Normalize destination path.
    dst = Path(output_path)
    # Validate source DOCX exists.
    if not src.exists() or not src.is_file():
        raise FileWriteError(f"Input DOCX not found: {src}")
    # Ensure destination folder exists.
    _ensure_parent_dir(dst)

    # Open source DOCX.
    try:
        doc = Document(str(src))
    # Wrap open/parsing errors.
    except Exception as e:
        raise FileWriteError(f"Failed to open DOCX: {src} ({e})") from e

    # Flatten relevant paragraphs from document parts.
    paragraphs = _collect_docx_paragraphs(doc)
    # Resolve transformation predicate (custom or default).
    predicate = should_transform or _should_transform_default
    # Counter for paragraphs visited.
    processed = 0
    # Counter for paragraphs changed.
    changed = 0
    # Total paragraph count for progress reporting.
    total = len(paragraphs)

    # Process each paragraph.
    for paragraph in paragraphs:
        # Increment progress counter.
        processed += 1
        # Emit progress callback when provided.
        if on_progress is not None:
            on_progress(processed, total)
        # Current paragraph text.
        original = paragraph.text or ""
        # Skip paragraphs rejected by predicate.
        if not predicate(original):
            continue
        # Compute transformed text.
        updated = transform(original)
        # Write back only if content changed.
        if updated != original:
            _set_paragraph_text_preserving_runs(paragraph, updated)
            changed += 1

    # Save transformed DOCX to destination.
    try:
        doc.save(str(dst))
    # Wrap save failures.
    except OSError as e:
        raise FileWriteError(f"Failed to write DOCX: {dst} ({e})") from e

    # Return (processed_count, changed_count).
    return processed, changed


# Write output based on extension (.docx uses DOCX writer, otherwise plain text).
def write_any_file(path: str | Path, text: str) -> None:
    # Normalize destination path.
    p = Path(path)
    # Compute lowercase extension.
    ext = p.suffix.lower()
    # Route DOCX outputs to DOCX writer.
    if ext == ".docx":
        write_docx_file(p, text)
        return
    # Fallback: write plain UTF-8 text.
    write_text_file(p, text)
